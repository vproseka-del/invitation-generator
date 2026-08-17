import io

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def text_to_docx_bytes(text):
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(14)
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(0)
    paragraph_format.space_before = Pt(0)
    paragraph_format.line_spacing = 1.15

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)

    paragraphs = text.split("\n")

    i = 0
    while i < len(paragraphs):
        line = paragraphs[i].strip()

        if not line:
            i += 1
            continue

        if _is_signature_block(line, paragraphs, i):
            i = _add_signature_block(doc, paragraphs, i)
            continue

        if _is_contacts_block(line):
            i = _add_contacts_block(doc, paragraphs, i)
            continue

        if _is_centered_line(line):
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        if _is_closing(line):
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _is_closing(line):
    return line.rstrip().endswith(",")


def _is_centered_line(line):
    return False


def _is_contacts_block(line):
    return line.lower().startswith("контактная информация")


def _is_signature_block(line, paragraphs, index):
    return (
        index > 0
        and line
        and not line.lower().startswith("контактная")
        and not line.lower().startswith("с уважением")
        and len(line.split()) <= 5
        and paragraphs[index - 1].strip().rstrip().endswith(",")
    )


def _add_contacts_block(doc, paragraphs, index):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run(paragraphs[index])
    run.bold = True

    for i in range(index + 1, len(paragraphs)):
        line = paragraphs[i].strip()
        if not line:
            continue
        contact_p = doc.add_paragraph(line)
        contact_p.paragraph_format.space_before = Pt(2)
        return i + 1
    return index + 1


def _add_signature_block(doc, paragraphs, index):
    for i in range(index, len(paragraphs)):
        line = paragraphs[i].strip()
        if not line:
            return i + 1
        if line.lower().startswith("контактная"):
            return _add_contacts_block(doc, paragraphs, i)
        p = doc.add_paragraph(line)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(0)
    return len(paragraphs)
